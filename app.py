# app.py
import streamlit as st
import pandas as pd
from openai import OpenAI
from docx import Document
from datetime import datetime
import io

# Page config
st.set_page_config(page_title="Excel Summarizer", page_icon="📊", layout="wide")

# Load API key from Streamlit secrets (secure method)
try:
    api_key = st.secrets["OPENAI_API_KEY"]
except:
    st.error("⚠️ API key not configured. Please contact administrator.")
    st.stop()

# Initialize OpenAI client with pre-loaded API key
client = OpenAI(api_key=api_key)

# Title
st.title("📊 Excel Spreadsheet Summarizer")
st.markdown("Upload Excel files and get AI-powered summaries")

# Main content
uploaded_files = st.file_uploader(
    "Choose Excel files", 
    type=['xlsx', 'xls'], 
    accept_multiple_files=True,
    help="Upload one or more Excel files to analyze"
)

def read_excel_file(uploaded_file):
    """Read an Excel file and return all sheets as a dictionary of DataFrames"""
    try:
        excel_file = pd.ExcelFile(uploaded_file)
        sheets_data = {}
        for sheet_name in excel_file.sheet_names:
            sheets_data[sheet_name] = pd.read_excel(uploaded_file, sheet_name=sheet_name)
        return sheets_data
    except Exception as e:
        st.error(f"Error reading file: {e}")
        return None

def prepare_data_for_summary(sheets_data):
    """Convert Excel data to text format for summarization"""
    text_content = ""
    for sheet_name, df in sheets_data.items():
        text_content += f"\n\n=== Sheet: {sheet_name} ===\n"
        text_content += f"Shape: {df.shape[0]} rows × {df.shape[1]} columns\n\n"
        text_content += f"Columns: {', '.join(df.columns.tolist())}\n\n"
        text_content += "Sample data:\n"
        text_content += df.head(10).to_string()
        text_content += "\n\n"
        if df.select_dtypes(include=['number']).shape[1] > 0:
            text_content += "Numeric column statistics:\n"
            text_content += df.describe().to_string()
    return text_content

def summarize_with_openai(content, filename):
    """Use OpenAI to generate a summary of the spreadsheet"""
    prompt = f"""Please analyze this Excel spreadsheet data from the file '{filename}' and provide a comprehensive summary including:
    
1. Overview of the data structure (sheets, columns, row counts)
2. Key insights and patterns in the data
3. Notable statistics or trends
4. Any data quality observations

Here's the spreadsheet data:

{content}
"""
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a data analyst expert who provides clear, concise summaries of spreadsheet data."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=2000
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"Error with OpenAI API: {e}")
        return None

def create_summary_document(summary, filename):
    """Create a Word document with the summary and return as bytes"""
    doc = Document()
    doc.add_heading(f'Spreadsheet Summary: {filename}', 0)
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph('')
    doc.add_heading('Summary', 1)
    doc.add_paragraph(summary)
    
    # Save to bytes buffer
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

# Process files when button is clicked
if uploaded_files:
    if st.button("🚀 Generate Summary", type="primary"):
        
        for uploaded_file in uploaded_files:
            with st.expander(f"📄 Processing: {uploaded_file.name}", expanded=True):
                # Read Excel file
                with st.spinner("Reading Excel file..."):
                    sheets_data = read_excel_file(uploaded_file)
                
                if sheets_data:
                    # Show preview
                    st.info(f"Found {len(sheets_data)} sheet(s)")
                    for sheet_name, df in sheets_data.items():
                        st.write(f"**{sheet_name}**: {df.shape[0]} rows × {df.shape[1]} columns")
                    
                    # Prepare data
                    content = prepare_data_for_summary(sheets_data)
                    
                    # Generate summary
                    with st.spinner("Generating AI summary..."):
                        summary = summarize_with_openai(content, uploaded_file.name)
                    
                    if summary:
                        # Display summary
                        st.success("✅ Summary generated!")
                        st.markdown("### Summary")
                        st.write(summary)
                        
                        # Create downloadable document
                        doc_buffer = create_summary_document(summary, uploaded_file.name)
                        output_filename = f"{uploaded_file.name.rsplit('.', 1)[0]}_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                        
                        st.download_button(
                            label="📥 Download Word Document",
                            data=doc_buffer,
                            file_name=output_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    
                    st.markdown("---")
else:
    st.info("👆 Upload Excel files to get started")

# Footer
st.markdown("---")
st.markdown("Built with Streamlit and OpenAI")